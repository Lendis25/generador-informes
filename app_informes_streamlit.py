
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import re
import tempfile

def calcular_promedio(df, run):
    if run not in df["RUN"].values:
        return ""
    notas = df[df["RUN"] == run].iloc[0].drop(labels=["RUN", "Nombre"], errors="ignore")
    notas = pd.to_numeric(notas, errors="coerce")
    return round(notas.dropna().mean(), 1) if not notas.dropna().empty else ""

def generar_informes_con_promedios(archivo_excel, archivo_plantilla, curso):
    datos = pd.read_excel(archivo_excel, sheet_name=None)
    estudiantes = datos["Lenguaje"][["RUN", "Nombre"]].dropna()

    asignaturas = {
        "Lenguaje y Comprensión": "Lenguaje",
        "I. Extranjero (Inglés)": "Inglés",
        "Ed. Matemática": "Matemática",
        "Est. Comp. De la Sociedad": "Historia y Geografía",
        "Est. y Comp. De la Naturaleza": "Ciencias Naturales",
        "Ed. Física": "Ed.Física"
    }

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(archivo_plantilla.getbuffer())
        plantilla_path = tmp.name

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        for _, row in estudiantes.iterrows():
            run = row["RUN"]
            nombre = row["Nombre"]
            doc = Document(plantilla_path)

            for p in doc.paragraphs:
                if "NOMBRE DEL ESTUDIANTE" in p.text:
                    p.text = f"NOMBRE DEL ESTUDIANTE: {nombre}"
                elif "CURSO:" in p.text:
                    p.text = f"CURSO: {curso}"

            # Llenar promedios en la tabla
            for table in doc.tables:
                for row in table.rows[1:]:  # Saltar encabezado
                    asignatura = row.cells[0].text.strip()
                    if asignatura in asignaturas:
                        hoja = asignaturas[asignatura]
                        promedio = calcular_promedio(datos[hoja], run)
                        row.cells[1].text = str(promedio)

            safe_name = re.sub(r"[^\w\s]", "", nombre).replace(" ", "_")
            doc_io = BytesIO()
            doc.save(doc_io)
            zipf.writestr(f"Informe_{safe_name}.docx", doc_io.getvalue())

    zip_buffer.seek(0)
    return zip_buffer

st.title("Generador de Informes con Promedios Automáticos")
st.write("Sube un archivo Excel con notas y una plantilla Word con tabla vacía.")

archivo_excel = st.file_uploader("Sube el archivo Excel (.xlsx)", type=["xlsx"])
archivo_word = st.file_uploader("Sube la plantilla Word (.docx)", type=["docx"])
curso = st.text_input("Nombre del curso (ej. 8° Básico)", value="8° Básico")

if archivo_excel and archivo_word and curso:
    if st.button("Generar Informes"):
        with st.spinner("Generando documentos Word con promedios..."):
            resultado = generar_informes_con_promedios(archivo_excel, archivo_word, curso)
            st.success("¡Informes generados con éxito!")
            st.download_button(
                label="Descargar informes (.zip)",
                data=resultado,
                file_name=f"Informes_{curso.replace(' ', '_')}.zip",
                mime="application/zip"
            )
